#!/usr/bin/env python
"""Non-interactive CLI port of 03_다중_에이전트_다중_에이전트로_이미지가_포함된_리포트_자동화해보기.ipynb.

Builds a report on a given topic using a LangGraph multi-agent pipeline:
  outline_generator -> contents_writer (Tavily search) -> image_generator (gpt-image-1)
  -> report_generator (python-docx), looping over sections until the report is complete.

Unlike the notebook (written for Google Colab), this reads secrets from a local
.env file and writes output under the current working directory instead of /content.

Usage:
    # Non-interactive (used by the /generate-report skill):
    python generate_report.py --topic "피지컬 AI" --sections 3

    # Interactive (run directly, prompts for topic/section count like the original notebook):
    python generate_report.py
"""
import argparse
import base64
import os
import uuid
from io import BytesIO
from typing import Annotated, Dict, List, Sequence, TypedDict

import requests
from dotenv import load_dotenv
from langchain_community.tools.tavily_search import TavilySearchResults
from langchain_core.messages import AIMessage, BaseMessage, HumanMessage
from langchain_core.output_parsers import JsonOutputParser
from langchain_core.prompts import PromptTemplate
from langchain_openai import ChatOpenAI
from langgraph.graph import END, START, StateGraph
from langgraph.graph.message import add_messages
from openai import OpenAI
from pydantic import Field, create_model


class State(TypedDict):
    messages: Annotated[Sequence[BaseMessage], add_messages]
    outline: Dict[str, str]
    current_section: int
    section_content: str
    section_image: str
    image_prompt: str
    total_sections: int
    full_report: List[Dict[str, str]]
    report_file: str


def create_outline_model(section_count: int):
    fields = {
        f"section{i}": (str, Field(description=f"Title for section {i}"))
        for i in range(1, section_count + 1)
    }
    return create_model("DynamicOutline", **fields)


def build_graph(llm: ChatOpenAI, search: TavilySearchResults, openai_client: OpenAI, output_dir: str):
    def outline_generator(state: State):
        outline_model = create_outline_model(state["total_sections"])
        outline_parser = JsonOutputParser(pydantic_object=outline_model)

        outline_prompt = PromptTemplate(
            template="""
            Create an outline for a detailed report with exactly {section_count} main sections.
            {format_instructions}
            The topic is: {topic}
            """,
            input_variables=["section_count", "topic"],
            partial_variables={"format_instructions": outline_parser.get_format_instructions()},
        )

        chain = outline_prompt | llm | outline_parser
        outline = chain.invoke(
            {
                "section_count": state["total_sections"],
                "topic": state["messages"][-1].content,
            }
        )
        return {"outline": outline}

    def contents_writer(state: State):
        if state["current_section"] > state["total_sections"]:
            return {"messages": [AIMessage(content="Report completed.")]}

        current_section_key = f"section{state['current_section']}"
        current_topic = state["outline"][current_section_key]
        search_results = search.invoke(current_topic)

        previous_sections_content = []
        for section in state.get("full_report", []):
            previous_sections_content.append(f"\n            {section['title']}\n            {section['content']}\n            ")
        previous_sections = "\n\n".join(previous_sections_content)

        section_prompt = PromptTemplate(
            template="""
            Write a detailed section for the topic: {topic}.

            Use the following search results for information: {search_results}

            Previous sections:
            {previous_sections}
            Write only the content for this section,
            do not include any image prompts or suggestions.
            Detailed statistics or information is needed,
            so you should include collected information from search result.""",
            input_variables=["topic", "search_results", "previous_sections"],
        )
        section_content = llm.invoke(
            section_prompt.format(
                topic=current_topic,
                search_results=search_results,
                previous_sections=previous_sections,
            )
        )

        return {
            "section_content": section_content.content,
            "current_section": state["current_section"],
        }

    def generate_image_gemini(prompt: str, api_key: str) -> bytes:
        """Generate an image with Gemini (gemini-3.1-flash-image) and return raw image bytes."""
        from google import genai
        from google.genai import types

        client = genai.Client(api_key=api_key)
        response = client.models.generate_content(
            model="gemini-3.1-flash-image",
            contents=[prompt],
            config=types.GenerateContentConfig(response_modalities=["TEXT", "IMAGE"]),
        )
        for part in response.candidates[0].content.parts:
            if part.inline_data and part.inline_data.data:
                return part.inline_data.data
        raise RuntimeError("Gemini response contained no image data")

    def generate_image_openai(prompt: str) -> bytes:
        """Generate an image with OpenAI's gpt-image-1 and return raw image bytes."""
        response = openai_client.images.generate(
            model="gpt-image-1",
            prompt=prompt,
            size="1024x1024",
            quality="medium",
            n=1,
        )
        return base64.b64decode(response.data[0].b64_json)

    def generate_image(prompt: str) -> str:
        """Generate an image and save it locally, returning the file path.

        Prefers Gemini (gemini-3.1-flash-image) when GEMINI_API_KEY/GOOGLE_API_KEY is set,
        falling back to OpenAI's gpt-image-1 if Gemini isn't configured or the call fails
        (e.g. a free-tier quota limit) — so a Gemini outage never breaks report generation.
        """
        image_bytes = None
        gemini_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
        if gemini_key:
            try:
                image_bytes = generate_image_gemini(prompt, gemini_key)
                print("Image generated with Gemini (gemini-3.1-flash-image).", flush=True)
            except Exception as exc:
                print(f"Gemini image generation failed ({exc}); falling back to OpenAI gpt-image-1.", flush=True)

        if image_bytes is None:
            image_bytes = generate_image_openai(prompt)
            print("Image generated with OpenAI (gpt-image-1).", flush=True)

        os.makedirs(output_dir, exist_ok=True)
        image_path = os.path.join(output_dir, f"{uuid.uuid4()}.png")
        with open(image_path, "wb") as f:
            f.write(image_bytes)
        return image_path

    def image_generator(state: State):
        prompt_template = PromptTemplate(
            template="""
            Based on the following section content, create a prompt for generating an infographic that represents this section.

            Section content:
            {section_content}

            Image generation prompt under 500 characters:
            """,
            input_variables=["section_content"],
        )
        image_prompt = llm.invoke(prompt_template.format(section_content=state["section_content"]))
        image_prompt_text = image_prompt.content if isinstance(image_prompt, AIMessage) else str(image_prompt)

        try:
            image_path = generate_image(image_prompt_text)
        except Exception as exc:  # keep the pipeline going even if image generation fails
            print(f"Image generation failed: {exc}", flush=True)
            image_path = "Image generation failed"

        current_section = {
            "title": state["outline"][f"section{state['current_section']}"],
            "content": state["section_content"],
            "image_url": image_path,
            "image_prompt": image_prompt_text,
        }
        updated_full_report = state.get("full_report", []) + [current_section]

        print(f"Completed section {state['current_section']} of {state['total_sections']}", flush=True)

        return {
            "image_prompt": image_prompt_text,
            "section_image": image_path,
            "current_section": state["current_section"] + 1,
            "full_report": updated_full_report,
        }

    def report_generator(state: State):
        from docx import Document
        from docx.shared import Inches

        doc = Document()
        doc.add_heading(f"Report: {state['messages'][0].content}", 0)

        for section in state["full_report"]:
            doc.add_heading(section["title"], level=1)
            doc.add_paragraph(section["content"])

            if section["image_url"] != "Image generation failed":
                try:
                    if section["image_url"].startswith("http"):
                        response = requests.get(section["image_url"])
                        image = BytesIO(response.content)
                    else:
                        image = section["image_url"]  # local file path
                    doc.add_picture(image, width=Inches(6))
                    doc.add_paragraph(f"Image prompt: {section['image_prompt']}")
                except Exception as e:
                    doc.add_paragraph(f"Failed to add image: {str(e)}")

            doc.add_page_break()

        safe_topic = "".join(c for c in state["messages"][0].content if c.isalnum() or c in " _-").strip()
        filename = os.path.join(output_dir, f"report_{safe_topic}.docx".replace(" ", "_"))
        doc.save(filename)

        return {
            "messages": [AIMessage(content=f"Report finalized and saved as {filename}.")],
            "report_file": filename,
        }

    def should_continue_writing(state: State):
        if state["current_section"] <= state["total_sections"]:
            return "write_section"
        return "finalize_report"

    graph_builder = StateGraph(State)
    graph_builder.add_node("outline_generator", outline_generator)
    graph_builder.add_node("contents_writer", contents_writer)
    graph_builder.add_node("image_generator", image_generator)
    graph_builder.add_node("report_generator", report_generator)

    graph_builder.add_edge(START, "outline_generator")
    graph_builder.add_edge("outline_generator", "contents_writer")
    graph_builder.add_edge("contents_writer", "image_generator")
    graph_builder.add_edge("report_generator", END)
    graph_builder.add_conditional_edges(
        "image_generator",
        should_continue_writing,
        {"write_section": "contents_writer", "finalize_report": "report_generator"},
    )

    return graph_builder.compile()


def main():
    parser = argparse.ArgumentParser(description="Generate an illustrated .docx report on a topic.")
    parser.add_argument("--topic", help="Report topic. Omit to be prompted for it interactively.")
    parser.add_argument("--sections", type=int, help="Number of report sections. Omit to be prompted for it.")
    parser.add_argument("--output-dir", default="generated_reports", help="Directory for the .docx and images")
    args = parser.parse_args()

    # Interactive mode: run `python generate_report.py` with no flags and it behaves like the
    # original notebook, asking for the topic and section count on stdin.
    topic = args.topic or input("보고서 주제를 입력하세요: ").strip()
    if not topic:
        raise SystemExit("주제를 입력해야 합니다.")

    if args.sections is not None:
        sections = args.sections
    else:
        raw = input("생성할 섹션의 수를 입력하세요 (기본값 3): ").strip()
        sections = int(raw) if raw else 3

    load_dotenv()
    for var in ("OPENAI_API_KEY", "TAVILY_API_KEY"):
        if not os.environ.get(var):
            raise SystemExit(f"Missing {var} in the environment/.env file.")

    llm = ChatOpenAI(model="gpt-4o-mini")
    search = TavilySearchResults(max_results=3)
    openai_client = OpenAI()

    graph = build_graph(llm, search, openai_client, args.output_dir)

    initial_state = {
        "messages": [HumanMessage(content=topic)],
        "total_sections": sections,
        "current_section": 1,
    }

    final_state = None
    for chunk in graph.stream(initial_state):
        final_state = chunk

    print("\n=== 보고서 생성 완료 ===")
    if final_state:
        print(final_state)


if __name__ == "__main__":
    main()
